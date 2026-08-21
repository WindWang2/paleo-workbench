from __future__ import annotations

from pathlib import Path

import pytest

from paleo_workbench.project.models import ResourceItem
from paleo_workbench.ui.pages.preview_provider import PreviewProvider
from paleo_workbench.ui.pages.preview_settings import PreviewSettings
from paleo_workbench.resources.preview_parsers.models import PreviewResult


@pytest.mark.parametrize("ext", ["mp4", "mov", "webm", "mkv", "avi"])
def test_video_preview_returns_media(tmp_path: Path, ext: str):
    path = tmp_path / f"clip.{ext}"
    path.write_bytes(b"\x00" * 64)
    resource = ResourceItem(name=f"clip.{ext}", path=str(path), type="video", format=ext, status="parsed")
    result = PreviewProvider().preview(resource)
    assert result.mode == "media"
    assert result.media_path == str(path)
    assert result.format == ext


def test_doc_preview_message(tmp_path: Path):
    path = tmp_path / "legacy.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0")
    resource = ResourceItem(name="legacy.doc", path=str(path), type="document", format="doc", status="parsed")
    result = PreviewProvider().preview(resource)
    assert result.mode == "message"
    assert result.message == "旧版二进制 .doc 不受支持，请另存为 .docx 后再预览"


def test_docx_preview_paragraphs_and_table_order(tmp_path: Path):
    from docx import Document

    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("第一段")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "a1"
    table.cell(0, 1).text = "b1"
    table.cell(1, 0).text = "a2"
    table.cell(1, 1).text = "b2"
    doc.add_paragraph("第二段")
    doc.save(str(path))

    resource = ResourceItem(name="sample.docx", path=str(path), type="document", format="docx", status="parsed")
    result = PreviewProvider().preview(resource)
    assert result.mode == "text"
    # paragraphs in order, table as TSV at position
    # Expect text contains both paragraphs and TSV lines
    assert "第一段" in result.text
    assert "第二段" in result.text
    # TSV: cells joined by tab, one line per row
    assert "a1\tb1" in result.text
    assert "a2\tb2" in result.text
    # Order: 第一段 before table before 第二段
    idx1 = result.text.index("第一段")
    idx_t1 = result.text.index("a1\tb1")
    idx2 = result.text.index("第二段")
    assert idx1 < idx_t1 < idx2


def test_docx_preview_only_paragraphs(tmp_path: Path):
    from docx import Document

    path = tmp_path / "only_text.docx"
    doc = Document()
    doc.add_paragraph("Hello")
    doc.add_paragraph("World")
    doc.save(str(path))
    resource = ResourceItem(name="only_text.docx", path=str(path), type="document", format="docx", status="parsed")
    result = PreviewProvider().preview(resource)
    assert result.mode == "text"
    assert "Hello" in result.text
    assert "World" in result.text


def test_docx_preview_truncation(tmp_path: Path):
    from docx import Document

    path = tmp_path / "large.docx"
    doc = Document()
    # Create many paragraphs to exceed limit (text_limit_kib 16 -> small)
    for i in range(5000):
        doc.add_paragraph(f"段落 {i} " + "x" * 100)
    doc.save(str(path))
    from dataclasses import replace

    settings = replace(PreviewSettings.defaults(), text_limit_kib=16)
    resource = ResourceItem(name="large.docx", path=str(path), type="document", format="docx", status="parsed")
    result = PreviewProvider(settings).preview(resource)
    assert result.mode == "text"
    assert result.truncated is True
    assert "仅显示前 16 KiB" in result.warning
